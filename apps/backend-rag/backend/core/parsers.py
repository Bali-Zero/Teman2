"""
ZANTARA RAG - Document Parsers
Extract text from PDF and EPUB files
"""

import asyncio
import logging
import os
from pathlib import Path

try:
    from PyPDF2 import PdfReader
except ImportError:
    from pypdf import PdfReader

try:
    from docx import Document
except ImportError:
    Document = None

import ebooklib
from bs4 import BeautifulSoup
from ebooklib import epub

logger = logging.getLogger(__name__)


class DocumentParseError(Exception):
    """Custom exception for document parsing errors"""

    pass


async def extract_text_from_pdf_async(file_path: str, use_ocr: bool = False) -> str:
    """
    Async version of PDF extraction for use in async contexts.
    """
    try:
        logger.info(f"Parsing PDF (async): {file_path}")

        reader = PdfReader(file_path)
        text_parts = []

        for page_num, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Error extracting page {page_num}: {e}")
                continue

        full_text = "\n\n".join(text_parts)

        # If no text found, try OCR (for scanned PDFs)
        if not full_text.strip() or use_ocr:
            logger.info("No text found in PDF, attempting async OCR extraction...")
            try:
                ocr_text = await extract_text_from_pdf_ocr_async(file_path)
                if ocr_text and ocr_text.strip():
                    logger.info(f"OCR extraction successful: {len(ocr_text)} characters")
                    return ocr_text
                if not full_text.strip():
                    raise DocumentParseError(
                        f"No text extracted from PDF (even with OCR/Vision): {file_path}"
                    )
            except DocumentParseError:
                raise
            except Exception as ocr_error:
                logger.warning(f"OCR extraction failed: {ocr_error}")
                if not full_text.strip():
                    raise DocumentParseError(f"No text extracted from PDF: {file_path}")

        if not full_text.strip():
            raise DocumentParseError(f"No text extracted from PDF: {file_path}")

        logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
        return full_text

    except DocumentParseError:
        raise
    except Exception as e:
        raise DocumentParseError(f"Failed to parse PDF {file_path}: {str(e)}") from e


def extract_text_from_pdf(file_path: str, use_ocr: bool = False) -> str:
    """
    Extract text content from PDF file.
    Falls back to OCR if no text is found (for scanned PDFs).

    Args:
        file_path: Path to PDF file
        use_ocr: Force OCR even if text extraction succeeds

    Returns:
        Extracted text as string

    Raises:
        DocumentParseError: If parsing fails
    """
    try:
        logger.info(f"Parsing PDF: {file_path}")

        reader = PdfReader(file_path)
        text_parts = []

        for page_num, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Error extracting page {page_num}: {e}")
                continue

        full_text = "\n\n".join(text_parts)

        # If no text found, try OCR (for scanned PDFs)
        if not full_text.strip() or use_ocr:
            logger.info("No text found in PDF, attempting OCR extraction...")
            try:
                ocr_text = extract_text_from_pdf_ocr(file_path)
                if ocr_text and ocr_text.strip():
                    logger.info(f"OCR extraction successful: {len(ocr_text)} characters")
                    return ocr_text
                if not full_text.strip():
                    # Last resort: try vision model (skip if already in async context)
                    logger.info("OCR failed, trying vision model as last resort...")
                    try:
                        import asyncio

                        from backend.services.multimodal.pdf_vision_service import PDFVisionService

                        vision_service = PDFVisionService()
                        with open(file_path, "rb") as f:
                            pdf_bytes = f.read()
                        # Check if we're in an async context
                        try:
                            asyncio.get_running_loop()
                            # We're in async context, can't use asyncio.run()
                            logger.warning("Cannot use vision model in async context - OCR skipped")
                        except RuntimeError:
                            # No running loop, safe to use asyncio.run()
                            vision_text = asyncio.run(vision_service.extract_text(pdf_bytes))
                            if vision_text and vision_text.strip():
                                logger.info(
                                    f"Vision extraction successful: {len(vision_text)} characters"
                                )
                                return vision_text
                    except Exception as vision_err:
                        logger.warning(f"Vision extraction failed: {vision_err}")

                    raise DocumentParseError(
                        f"No text extracted from PDF (even with OCR/Vision): {file_path}"
                    )
            except DocumentParseError:
                raise
            except Exception as ocr_error:
                logger.warning(f"OCR extraction failed: {ocr_error}")
                if not full_text.strip():
                    raise DocumentParseError(f"No text extracted from PDF: {file_path}")

        if not full_text.strip():
            raise DocumentParseError(f"No text extracted from PDF: {file_path}")

        logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
        return full_text

    except DocumentParseError:
        raise
    except Exception as e:
        raise DocumentParseError(f"Failed to parse PDF {file_path}: {str(e)}") from e


async def extract_text_from_pdf_ocr_async(file_path: str) -> str:
    """
    Async version of OCR extraction for use in async contexts.
    Uses Google Gemini Vision to extract text from scanned PDF pages.
    Processes each page individually for best OCR quality.
    """
    try:
        import fitz  # PyMuPDF

        from backend.services.multimodal.pdf_vision_service import PDFVisionService

        vision_service = PDFVisionService()
        if not vision_service._available:
            logger.warning("Vision service not available for OCR")
            return ""

        # Open PDF and get page count
        doc = fitz.open(file_path)
        total_pages = len(doc)
        doc.close()

        logger.info(f"Processing {total_pages} pages with Gemini Vision OCR...")

        # OCR prompt optimized for text extraction
        ocr_prompt = """Extract all text from this page.
Preserve the original formatting, line breaks, paragraphs, and structure.
Return only the extracted text, no descriptions or explanations.
If the page contains tables, preserve the table structure.
If the page is blank or contains no text, return an empty string."""

        text_parts = []
        for page_num in range(1, total_pages + 1):
            try:
                logger.info(f"OCR processing page {page_num}/{total_pages}...")
                page_text = await vision_service.analyze_page(
                    pdf_path=file_path, page_number=page_num, prompt=ocr_prompt, is_drive_file=False
                )
                if page_text and page_text.strip():
                    text_parts.append(page_text)
                    logger.info(f"✅ Page {page_num}: extracted {len(page_text)} characters")
                else:
                    logger.warning(f"⚠️ Page {page_num}: no text extracted")

                # Rate limiting for Gemini Vision (Free Tier: 15 RPM)
                # 4s delay = 15 RPM
                await asyncio.sleep(4.0)
            except Exception as page_error:
                logger.warning(f"Error processing page {page_num}: {page_error}")
                continue

        full_text = "\n\n".join(text_parts)

        if full_text.strip():
            logger.info(
                f"✅ Gemini Vision OCR successful: {len(full_text)} characters from {len(text_parts)} pages"
            )
            return full_text
        logger.warning("No text extracted from any page")
        return ""

    except Exception as e:
        logger.error(f"Vision OCR extraction failed: {e}")
        return ""


def extract_text_from_pdf_ocr(file_path: str) -> str:
    """
    Extract text from scanned PDF using OCR (PyMuPDF with OCR support).

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text as string
    """
    try:
        import fitz  # PyMuPDF

        logger.info(f"Attempting OCR extraction from PDF: {file_path}")
        doc = fitz.open(file_path)
        text_parts = []

        for page_num, page in enumerate(doc, 1):
            try:
                # Try OCR if available (requires PyMuPDF with OCR support)
                # PyMuPDF can use Tesseract if available
                text = page.get_text("text")

                # If no text, try OCR mode
                if not text.strip():
                    # PyMuPDF OCR requires Tesseract
                    # Fallback: use page.get_text() with OCR flag if available
                    try:
                        # Try with OCR flag (if PyMuPDF was compiled with OCR support)
                        text = page.get_text("text", flags=11)  # OCR flag
                    except Exception:
                        # If OCR not available, try alternative extraction
                        logger.warning(
                            f"OCR not available for page {page_num}, trying alternative extraction"
                        )
                        # Convert page to image and use vision model as fallback
                        page.get_pixmap()
                        # For now, return empty - can be enhanced with vision model
                        text = ""

                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Error extracting page {page_num} with OCR: {e}")
                continue

        doc.close()
        return "\n\n".join(text_parts)

    except ImportError:
        logger.warning("PyMuPDF (fitz) not available for OCR extraction")
        return ""
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text as string

    Raises:
        DocumentParseError: If parsing fails
    """
    if Document is None:
        raise DocumentParseError("python-docx not installed. Install with: pip install python-docx")

    try:
        logger.info(f"Parsing DOCX: {file_path}")

        doc = Document(file_path)
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            raise DocumentParseError(f"No text extracted from DOCX: {file_path}")

        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        return full_text

    except Exception as e:
        raise DocumentParseError(f"Failed to parse DOCX {file_path}: {str(e)}") from e


def extract_text_from_epub(file_path: str) -> str:
    """
    Extract text content from EPUB file.

    Args:
        file_path: Path to EPUB file

    Returns:
        Extracted text as string

    Raises:
        DocumentParseError: If parsing fails
    """
    try:
        logger.info(f"Parsing EPUB: {file_path}")

        book = epub.read_epub(file_path)
        text_parts = []

        # Extract text from each chapter
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                try:
                    content = item.get_content().decode("utf-8")
                    soup = BeautifulSoup(content, "html.parser")
                    text = soup.get_text(separator="\n", strip=True)

                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting chapter: {e}")
                    continue

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            raise DocumentParseError(f"No text extracted from EPUB: {file_path}")

        logger.info(f"Successfully extracted {len(full_text)} characters from EPUB")
        return full_text

    except Exception as e:
        raise DocumentParseError(f"Failed to parse EPUB {file_path}: {str(e)}") from e


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text content from TXT file.

    Args:
        file_path: Path to TXT file

    Returns:
        Extracted text as string

    Raises:
        DocumentParseError: If parsing fails
    """
    try:
        logger.info(f"Reading TXT: {file_path}")
        with open(file_path, encoding="utf-8") as f:
            text = f.read()

        if not text.strip():
            raise DocumentParseError(f"No text extracted from TXT: {file_path}")

        logger.info(f"Successfully extracted {len(text)} characters from TXT")
        return text

    except Exception as e:
        raise DocumentParseError(f"Failed to read TXT {file_path}: {str(e)}") from e


def auto_detect_and_parse(file_path: str, use_ocr: bool = False) -> str:
    """
    Auto-detect file type and parse accordingly.

    Args:
        file_path: Path to document file
        use_ocr: Force OCR for PDFs (for scanned documents)

    Returns:
        Extracted text as string

    Raises:
        DocumentParseError: If file type unsupported or parsing fails
    """
    if not os.path.exists(file_path):
        raise DocumentParseError(f"File not found: {file_path}")

    file_ext = Path(file_path).suffix.lower()

    if file_ext == ".pdf":
        return extract_text_from_pdf(file_path, use_ocr=use_ocr)
    if file_ext == ".docx":
        return extract_text_from_docx(file_path)
    if file_ext == ".epub":
        return extract_text_from_epub(file_path)
    if file_ext == ".txt":
        return extract_text_from_txt(file_path)
    if file_ext == ".md":
        return extract_text_from_txt(file_path)  # Markdown is plain text
    raise DocumentParseError(
        f"Unsupported file type: {file_ext}. Supported formats: .pdf, .docx, .epub, .txt, .md"
    )


def get_document_info(file_path: str) -> dict:
    """
    Extract basic metadata from document.

    Args:
        file_path: Path to document

    Returns:
        Dictionary with document metadata
    """
    info = {
        "file_name": Path(file_path).name,
        "file_size_mb": os.path.getsize(file_path) / (1024 * 1024),
        "file_type": Path(file_path).suffix.lower(),
        "file_path": file_path,
    }

    try:
        if info["file_type"] == ".pdf":
            reader = PdfReader(file_path)
            info["pages"] = len(reader.pages)

            # Try to get PDF metadata
            if reader.metadata:
                info["title"] = reader.metadata.get("/Title", "")
                info["author"] = reader.metadata.get("/Author", "")

        elif info["file_type"] == ".docx":
            if Document is None:
                logger.warning("python-docx not available for DOCX metadata extraction")
            else:
                doc = Document(file_path)
                # Try to get document properties
                props = doc.core_properties
                info["title"] = props.title or ""
                info["author"] = props.author or ""
                info["pages"] = len(doc.paragraphs)  # Approximate page count

        elif info["file_type"] == ".epub":
            book = epub.read_epub(file_path)
            info["title"] = (
                book.get_metadata("DC", "title")[0][0] if book.get_metadata("DC", "title") else ""
            )
            info["author"] = (
                book.get_metadata("DC", "creator")[0][0]
                if book.get_metadata("DC", "creator")
                else ""
            )

    except Exception as e:
        logger.warning(f"Could not extract metadata from {file_path}: {e}")

    return info
