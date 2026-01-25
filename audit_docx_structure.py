
import docx
import sys
from pathlib import Path

FILE_PATH = Path("/Users/antonellosiano/Desktop/nuzantara/2.3 Lampiran I.C PP Nomor 28 Tahun 2025 (I.C.1-182) (1).docx")

def main():
    if not FILE_PATH.exists():
        print(f"File not found: {FILE_PATH}")
        return

    print(f"Analyzing: {FILE_PATH.name}")
    try:
        doc = docx.Document(FILE_PATH)
        print(f"Paragraphs: {len(doc.paragraphs)}")
        print(f"Tables: {len(doc.tables)}")

        if not doc.tables:
            print("❌ No tables found!")
            return

        first_table = doc.tables[0]
        rows = first_table.rows
        print(f"First Table Rows: {len(rows)}")
        
        # Check col count
        if len(rows) > 0:
            print(f"Columns in first row: {len(rows[0].cells)}")

        print("-" * 40)
        print("First 5 Rows Preview:")
        for i, row in enumerate(rows[:5]):
            cells = [c.text.strip() for c in row.cells]
            print(f"Row {i}: {cells}")
        print("-" * 40)

    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    main()
