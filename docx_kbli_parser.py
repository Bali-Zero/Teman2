
import docx
import json
import re
from pathlib import Path
from typing import List, Dict, Optional

# Configuration
INPUT_FILE = Path("/Users/antonellosiano/Desktop/nuzantara/2.3 Lampiran I.C PP Nomor 28 Tahun 2025 (I.C.1-182) (1).docx")
OUTPUT_JSON = Path("/Users/antonellosiano/Desktop/nuzantara/KBLI_Lampiran_IC_Masterpiece.json")

COLUMN_NAMES = [
    "no",
    "kode_kbli",
    "judul_kbli",
    "ruang_lingkup",
    "skala_usaha",
    "tingkat_risiko",
    "perizinan_berusaha",
    "persyaratan",
    "jangka_waktu_penerbitan",
    "kewajiban",
    "pb_uu_ikm",
    "parameter",
    "kewenangan",
]

def clean_text(text: str) -> str:
    if not text:
        return ""
    # Standardize
    text = text.replace("\u200b", "") # Remove zero-width space
    
    # Fix hyphenation (e.g. "Kehu-\ntanan" -> "Kehutanan")
    # Regex: word-\nword -> wordword
    text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)
    text = re.sub(r"(\w+)-\s+(\w+)", r"\1\2", text)
    
    return text.strip()

def process_docx(file_path: Path) -> List[Dict]:
    doc = docx.Document(file_path)
    records = []
    
    total_tables = len(doc.tables)
    print(f"Processing {total_tables} tables...")
    
    last_valid_kbli = None
    
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = row.cells
            
            if len(cells) < 13:
                continue

            cleaned_cells = [clean_text(c.text) for c in cells]
            
            # 1. SKIP GARBAGE / HEADERS
            c1 = cleaned_cells[1] # Kode
            c2 = cleaned_cells[2] # Judul
            
            # Header signatures
            if "kode" in c1.lower() and "kbli" in c1.lower():
                continue
            if "3)" in c2 or "131" in c2: # Column numbering garbage
                continue
            if c1 == "121" or c1 == "(2)": # Distinct garbage signature
                continue
            
            # 2. IDENTIFY KBLI
            current_kbli = None
            
            # Case A: Cell has 5 digits
            if c1 and len(c1) >= 5 and c1[0].isdigit():
                # Check if it's a real code (sometimes 121 passes IsDigit check if sloppy)
                # Regex for 5 digit info
                match = re.search(r"\b\d{5}\b", c1)
                if match:
                    current_kbli = match.group(0)
            
            # Case B: Implicit Continuation (Empty code, but has data)
            # Use last_valid_kbli if we have relevant data in Scope (3), Scale (4), or Risk (5)
            # But only if we are sure it's not a garbage row or empty row
            has_data = any(cleaned_cells[3:]) # Has content in scope, scale, risk etc
            if not current_kbli and not c1 and has_data and last_valid_kbli:
                 current_kbli = last_valid_kbli
                 # If implied, we might need to verify if Title is empty too (it should be)
            
            if not current_kbli:
                continue
                
            last_valid_kbli = current_kbli
            
            # 3. EXTRACTION
            record = {}
            # Map clean cells to columns
            # But ensure Kode KBLI is the cleaned/detected one
            record["kode_kbli"] = current_kbli
            
            # For other columns, blindly map
            for i, col_name in enumerate(COLUMN_NAMES):
                if col_name == "kode_kbli":
                    continue
                    
                if i < len(cleaned_cells):
                    val = cleaned_cells[i]
                    record[col_name] = val
                else:
                    record[col_name] = ""
            
            record["source_table"] = t_idx
            records.append(record)
            
    return records

def main():
    if not INPUT_FILE.exists():
        print(f"File not found: {INPUT_FILE}")
        return

    print("🚀 Nuzantara DOCX KBLI Parser")
    print(f"Target: {INPUT_FILE.name}")
    
    data = process_docx(INPUT_FILE)
    
    print(f"✅ Extracted {len(data)} records")
    
    # Save
    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump({
            "meta": {
                "source": INPUT_FILE.name,
                "count": len(data),
                "type": "Masterpiece DOCX Extraction"
            },
            "data": data
        }, f, indent=2, ensure_ascii=False)
        
    print(f"Saved to: {OUTPUT_JSON}")

if __name__ == "__main__":
    main()
